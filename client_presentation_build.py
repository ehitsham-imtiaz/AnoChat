from __future__ import annotations

import time
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from playwright.sync_api import sync_playwright


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "client_presentation"
SHOT_DIR = OUT / "screenshots"
DOCX = OUT / "AnoChat_Client_Presentation.docx"
APP_URL = "http://127.0.0.1:8000/frontend/index.html"
CHROME = r"C:\Program Files\Google\Chrome\Application\chrome.exe"


def wait_for_app(page):
    page.wait_for_timeout(1200)
    page.wait_for_load_state("networkidle", timeout=15000)


def screenshot(page, name):
    path = SHOT_DIR / f"{name}.png"
    page.screenshot(path=str(path), full_page=False)
    return path


def capture_screenshots():
    SHOT_DIR.mkdir(parents=True, exist_ok=True)
    shots = {}
    with sync_playwright() as p:
        browser = p.chromium.launch(executable_path=CHROME, headless=True)
        page = browser.new_page(viewport={"width": 1440, "height": 900}, device_scale_factor=1)
        page.route("**/static/env.js", lambda route: route.fulfill(
            status=200,
            content_type="application/javascript",
            body='(function(){ window.API_BASE = ""; })();',
        ))
        page.goto(APP_URL, wait_until="domcontentloaded", timeout=30000)
        page.evaluate("localStorage.clear()")
        page.reload(wait_until="domcontentloaded")
        wait_for_app(page)
        shots["login"] = screenshot(page, "01_login")

        page.fill('input[name="login"]', "admin@example.com")
        page.fill('input[name="password"]', "Admin123!")
        page.click('button:has-text("Log in to portal")')
        page.wait_for_selector(".main-shell", timeout=30000)
        wait_for_app(page)
        shots["dashboard"] = screenshot(page, "02_dashboard")

        for key, label, fname in [
            ("projects", "Projects", "03_projects"),
            ("chatters", "Chatter", "04_chatter"),
            ("monitoring", "Monitoring", "05_monitoring"),
            ("users", "Users & Roles", "06_users_roles"),
        ]:
            page.click(f'.nav-link:has-text("{label}")')
            page.wait_for_timeout(1800)
            shots[key] = screenshot(page, fname)

        page.click('.nav-link:has-text("Projects")')
        page.wait_for_timeout(900)
        if page.locator('button:has-text("New Project")').count():
            page.click('button:has-text("New Project")')
        elif page.locator('button:has-text("Create Project")').count():
            page.click('button:has-text("Create Project")')
        page.wait_for_timeout(900)
        shots["project_modal"] = screenshot(page, "07_create_project_modal")

        browser.close()
    return shots


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)


def add_section_header(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(15, 23, 42)
    if subtitle:
        p2 = doc.add_paragraph(subtitle)
        p2.paragraph_format.space_after = Pt(8)
        p2.runs[0].font.size = Pt(9.5)
        p2.runs[0].font.color.rgb = RGBColor(71, 85, 105)


def add_image_block(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.35))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(8.5)
    cap.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(9.7)
        r.font.color.rgb = RGBColor(30, 41, 59)


def build_doc(shots):
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(6)
    tr = title.add_run("AnoChat Workspace")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Client Presentation Document")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(37, 99, 235)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_before = Pt(14)
    intro.paragraph_format.space_after = Pt(40)
    ir = intro.add_run(
        "A modern standalone SaaS workspace for project management, team chatter, user roles, monitoring, and operational visibility."
    )
    ir.font.size = Pt(11)
    ir.font.color.rgb = RGBColor(71, 85, 105)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    labels = [("Frontend", "Vercel static workspace"), ("Backend", "FastAPI service"), ("Database", "PostgreSQL")]
    for idx, (head, body) in enumerate(labels):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "EEF4FF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(head + "\n")
        run.bold = True
        run.font.color.rgb = RGBColor(29, 78, 216)
        run.font.size = Pt(10)
        run2 = p.add_run(body)
        run2.font.color.rgb = RGBColor(71, 85, 105)
        run2.font.size = Pt(8.5)

    doc.add_page_break()

    add_section_header(doc, "1. Executive Overview", "What AnoChat gives the client")
    add_bullets(doc, [
        "A polished login experience and protected workspace access.",
        "Role-based access: admin users manage the complete system, while other roles see assigned work and chatter.",
        "Project records automatically connect to team conversation spaces.",
        "Chatter supports attachments, mentions, deleted-message handling, and contact-detail masking.",
        "Monitoring gives a clean audit trail for activity across the workspace.",
    ])

    add_section_header(doc, "2. Login Experience", "Secure entry point for all workspace users")
    add_image_block(doc, shots["login"], "Figure 1: Premium SaaS login page with AnoChat branding and role positioning.")
    add_bullets(doc, [
        "The login page uses a split-screen SaaS layout with clear brand positioning.",
        "Credentials are no longer prefilled, which is safer for production use.",
        "The frontend is configured to communicate with the deployed backend gateway.",
    ])

    doc.add_page_break()
    add_section_header(doc, "3. Dashboard", "A concise workspace command center")
    add_image_block(doc, shots["dashboard"], "Figure 2: Dashboard showing workspace health, quick actions, recent projects, and active chatter.")
    add_bullets(doc, [
        "The dashboard summarizes projects, chatters, users, and activity depending on role access.",
        "Quick action cards guide the user directly to core workflows.",
        "The layout is responsive and optimized for scanning rather than heavy admin clutter.",
    ])

    add_section_header(doc, "4. Projects", "Project management with linked chatter spaces")
    add_image_block(doc, shots["projects"], "Figure 3: Projects section with search, filtering, status, priority, ownership, and actions.")
    add_bullets(doc, [
        "Projects support creation, editing, deletion, status, priority, customer, and member assignment.",
        "Creating a project also creates the related chatter so teams can immediately collaborate.",
        "Archived/deleted project flows are tied to corresponding chatter cleanup behavior.",
    ])

    doc.add_page_break()
    add_section_header(doc, "5. Create Project Modal", "Modern form flow for new workspaces")
    add_image_block(doc, shots["project_modal"], "Figure 4: Create Project modal with project fields, customer selection, and member assignment.")
    add_bullets(doc, [
        "The form keeps project creation simple while preserving important metadata.",
        "Members can be assigned through a modern dropdown experience.",
        "The modal is responsive and uses clean validation-focused input styling.",
    ])

    add_section_header(doc, "6. Chatter", "Team conversation workspace")
    add_image_block(doc, shots["chatters"], "Figure 5: Chatter section with conversation list, message composer, attachments, and group context.")
    add_bullets(doc, [
        "Chatter supports messages, file attachments, image previews, and group/member context.",
        "Regex-based contact protection masks sensitive contact data for non-admin messages.",
        "Admin visibility rules preserve governance while protecting normal users from exposed contact data.",
        "The screen uses fixed chat layout behavior so the conversation area remains usable.",
    ])

    doc.add_page_break()
    add_section_header(doc, "7. Monitoring", "Audit events and workspace activity")
    add_image_block(doc, shots["monitoring"], "Figure 6: Monitoring screen with metrics, activity search, filters, and audit timeline.")
    add_bullets(doc, [
        "Monitoring provides visibility into login, project, chatter, message, user, and attachment activity.",
        "Search and type filters allow the admin to quickly inspect relevant audit events.",
        "The activity log was redesigned into a cleaner, minimal layout for client readability.",
    ])

    add_section_header(doc, "8. Users & Roles", "Access management for the workspace")
    add_image_block(doc, shots["users"], "Figure 7: Users & Roles screen with role badges, status indicators, and user actions.")
    add_bullets(doc, [
        "Admins can create, edit, deactivate, delete, and update user roles.",
        "Users can update their own password and presence status.",
        "Role-based navigation ensures non-admin users only see relevant assigned work and chatter.",
    ])

    doc.add_page_break()
    add_section_header(doc, "9. Deployment Summary", "Current production-oriented setup")
    dep = doc.add_table(rows=1, cols=2)
    dep.rows[0].cells[0].text = "Layer"
    dep.rows[0].cells[1].text = "Implementation"
    for cell in dep.rows[0].cells:
        set_cell_shading(cell, "1E3A8A")
        set_cell_text_color(cell, "FFFFFF")
    for layer, implementation in [
        ("Frontend", "Vercel static deployment from the /frontend directory."),
        ("Backend", "FastAPI service exposed through the Cloudflare gateway route."),
        ("Database", "PostgreSQL database for users, projects, chatters, messages, activity logs, and role data."),
        ("Attachments", "Stored through backend upload handling; persistent storage should be kept on the backend host."),
        ("Security", "JWT auth, role-based access, message masking, deleted-message visibility rules, and audit logging."),
    ]:
        row = dep.add_row()
        row.cells[0].text = layer
        row.cells[1].text = implementation

    for row in dep.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    add_section_header(doc, "10. Client Talking Points", "Suggested points for the presentation")
    add_bullets(doc, [
        "AnoChat is now independent from Odoo and runs as a standalone workspace.",
        "The UI has been redesigned into a modern SaaS experience suitable for production demos.",
        "Project creation, chatter, user management, monitoring, attachments, and role access are functional.",
        "The platform is ready for client review, with deployment split between Vercel frontend and backend API service.",
    ])

    doc.save(DOCX)
    return DOCX


def main():
    shots = capture_screenshots()
    docx_path = build_doc(shots)
    print(docx_path)


if __name__ == "__main__":
    main()
